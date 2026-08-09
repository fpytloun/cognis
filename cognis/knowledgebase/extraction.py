"""Artifact text extraction for knowledgebase indexing."""

from __future__ import annotations

import asyncio
import contextlib
import importlib.util
import io
import multiprocessing
import os
import pickle
import re
import tempfile
import zipfile
from dataclasses import dataclass, field
from typing import Any

import yaml


@dataclass(slots=True)
class SourceSpan:
    text: str
    locator: dict[str, Any] = field(default_factory=dict)


@dataclass(slots=True)
class ExtractedDocument:
    spans: list[SourceSpan]
    extraction_method: str
    metadata: dict[str, Any] = field(default_factory=dict)
    diagnostics: dict[str, Any] = field(default_factory=dict)


_TEXT_MIME_PREFIXES = ("text/",)
_TEXT_MIME_TYPES = {
    "application/json",
    "application/xml",
    "application/yaml",
    "application/x-yaml",
    "text/vtt",
    "application/x-subrip",
}
SUPPORTED_MIME_TYPES = tuple(
    sorted(
        {
            *_TEXT_MIME_TYPES,
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/plain",
            "text/markdown",
            "text/csv",
            "text/html",
        }
    )
)
SUPPORTED_EXTENSIONS = (
    ".csv",
    ".docx",
    ".html",
    ".json",
    ".md",
    ".pdf",
    ".srt",
    ".txt",
    ".vtt",
    ".xml",
    ".yaml",
    ".yml",
)


def available_supported_types() -> tuple[list[str], list[str]]:
    mime_types = list(SUPPORTED_MIME_TYPES)
    extensions = list(SUPPORTED_EXTENSIONS)
    if importlib.util.find_spec("docx") is None:
        mime_types.remove("application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        extensions.remove(".docx")
    if importlib.util.find_spec("pypdf") is None:
        mime_types.remove("application/pdf")
        extensions.remove(".pdf")
    return mime_types, extensions


def supports_artifact_type(*, filename: str, mime_type: str) -> bool:
    mime_types, extensions = available_supported_types()
    lower_name = filename.lower()
    return (
        lower_name.endswith(tuple(extensions))
        or mime_type.startswith(_TEXT_MIME_PREFIXES)
        or mime_type in mime_types
    )


_MAX_EXTRACTED_CHARACTERS = 20_000_000
_MAX_SOURCE_SPANS = 100_000
_MAX_PDF_PAGES = 2_000
_MAX_DOCX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
_MAX_DOCX_ARCHIVE_ENTRIES = 10_000
_MAX_DOCX_ENTRY_BYTES = 32 * 1024 * 1024
_MAX_DOCX_COMPRESSION_RATIO = 200
_MAX_FRONTMATTER_BYTES = 32 * 1024
_MAX_FRONTMATTER_KEYS = 64
_MAX_FRONTMATTER_DEPTH = 4
_MAX_FRONTMATTER_LIST_ITEMS = 100
_heavy_extraction_semaphore = asyncio.Semaphore(2)


class KnowledgebaseExtractionLimitExceeded(RuntimeError):
    pass


class KnowledgebaseExtractionTimeout(RuntimeError):
    pass


class KnowledgebaseMetadataEnvelopeError(RuntimeError):
    pass


class _NoAliasSafeLoader(yaml.SafeLoader):
    def compose_node(self, parent: Any, index: Any) -> Any:
        if self.check_event(yaml.AliasEvent):
            raise KnowledgebaseMetadataEnvelopeError("YAML aliases are not allowed")
        return super().compose_node(parent, index)


def _validate_metadata_value(value: Any, *, depth: int = 0) -> None:
    if depth > _MAX_FRONTMATTER_DEPTH:
        raise KnowledgebaseMetadataEnvelopeError("frontmatter exceeds nesting limit")
    if isinstance(value, dict):
        if len(value) > _MAX_FRONTMATTER_KEYS:
            raise KnowledgebaseMetadataEnvelopeError("frontmatter has too many keys")
        for key, item in value.items():
            if not isinstance(key, str) or not key or len(key) > 128:
                raise KnowledgebaseMetadataEnvelopeError("frontmatter keys must be bounded strings")
            _validate_metadata_value(item, depth=depth + 1)
    elif isinstance(value, list):
        if len(value) > _MAX_FRONTMATTER_LIST_ITEMS:
            raise KnowledgebaseMetadataEnvelopeError("frontmatter list exceeds item limit")
        for item in value:
            _validate_metadata_value(item, depth=depth + 1)
    elif value is not None and not isinstance(value, str | int | float | bool):
        raise KnowledgebaseMetadataEnvelopeError("frontmatter contains an unsupported value")


def _extract_markdown_frontmatter(text: str) -> tuple[str, dict[str, Any]]:
    if not text.startswith("---\n") and not text.startswith("---\r\n"):
        return text, {}
    match = re.search(r"\r?\n---\r?\n", text[4 : _MAX_FRONTMATTER_BYTES + 8])
    if match is None:
        raise KnowledgebaseMetadataEnvelopeError("frontmatter closing delimiter is missing")
    delimiter_start = 4 + match.start()
    body_start = 4 + match.end()
    envelope = text[4:delimiter_start]
    if len(envelope.encode("utf-8")) > _MAX_FRONTMATTER_BYTES:
        raise KnowledgebaseMetadataEnvelopeError("frontmatter exceeds size limit")
    try:
        parsed = yaml.load(envelope, Loader=_NoAliasSafeLoader)
    except yaml.YAMLError as exc:
        raise KnowledgebaseMetadataEnvelopeError("frontmatter is malformed") from exc
    if parsed is None:
        parsed = {}
    if not isinstance(parsed, dict):
        raise KnowledgebaseMetadataEnvelopeError("frontmatter root must be a mapping")
    _validate_metadata_value(parsed)
    return text[body_start:], parsed


def _is_heavy_type(*, filename: str, mime_type: str) -> bool:
    lower_name = filename.lower()
    return (
        lower_name.endswith((".pdf", ".docx"))
        or mime_type == "application/pdf"
        or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def _extract_document_worker(
    content: bytes, filename: str, mime_type: str, result_path: str
) -> None:
    with contextlib.suppress(ImportError, ValueError, OSError):
        import resource

        resource.setrlimit(resource.RLIMIT_CPU, (30, 30))
        memory_limit = 512 * 1024 * 1024
        resource.setrlimit(resource.RLIMIT_AS, (memory_limit, memory_limit))
    try:
        result = (
            "ok",
            extract_artifact_bytes(content, filename=filename, mime_type=mime_type),
        )
    except BaseException:
        result = ("error", None)
    with open(result_path, "wb") as output:
        pickle.dump(result, output)


async def _stop_extraction_process(process: Any) -> None:
    if not process.is_alive():
        return
    process.terminate()
    await asyncio.to_thread(process.join, 5)
    if process.is_alive() and hasattr(process, "kill"):
        process.kill()
        await asyncio.to_thread(process.join, 5)


async def extract_artifact_bytes_bounded(
    content: bytes,
    *,
    filename: str,
    mime_type: str,
    timeout_seconds: int = 30,
) -> ExtractedDocument:
    """Extract PDF/DOCX content in a small concurrency-limited process boundary."""

    if not _is_heavy_type(filename=filename, mime_type=mime_type):
        return extract_artifact_bytes(content, filename=filename, mime_type=mime_type)
    async with _heavy_extraction_semaphore:
        context = multiprocessing.get_context("spawn")
        descriptor, result_path = tempfile.mkstemp(prefix="cognis-kb-extract-")
        os.close(descriptor)
        process = context.Process(
            target=_extract_document_worker,
            args=(content, filename, mime_type, result_path),
            daemon=True,
        )
        started = False
        try:
            process.start()
            started = True
            await asyncio.to_thread(process.join, timeout_seconds)
            if process.is_alive():
                await _stop_extraction_process(process)
                raise KnowledgebaseExtractionTimeout("document extraction timed out")
            try:
                with open(result_path, "rb") as result_file:
                    status, document = pickle.load(result_file)
            except (OSError, EOFError, pickle.PickleError) as exc:
                raise RuntimeError("document extraction failed") from exc
            if status != "ok" or not isinstance(document, ExtractedDocument):
                raise RuntimeError("document extraction failed")
            return document
        finally:
            if started:
                await _stop_extraction_process(process)
            with contextlib.suppress(OSError):
                os.unlink(result_path)


def _timestamp_ms(raw: str) -> int:
    parts = raw.replace(",", ".").split(":")
    hours, minutes, seconds = int(parts[0]), int(parts[1]), float(parts[2])
    return int(((hours * 60 + minutes) * 60 + seconds) * 1000)


def _extract_transcript(text: str) -> ExtractedDocument:
    spans: list[SourceSpan] = []
    blocks = re.split(r"\n\s*\n", text.strip())
    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        time_line = next((line for line in lines if "-->" in line), None)
        if time_line is None:
            continue
        start_raw, end_raw = [part.strip().split()[0] for part in time_line.split("-->", 1)]
        payload = [
            line for line in lines if line != time_line and not line.isdigit() and line != "WEBVTT"
        ]
        if not payload:
            continue
        spans.append(
            SourceSpan(
                text=" ".join(payload),
                locator={
                    "timestamp_start_ms": _timestamp_ms(start_raw),
                    "timestamp_end_ms": _timestamp_ms(end_raw),
                },
            )
        )
    return ExtractedDocument(spans=spans, extraction_method="transcript")


def extract_artifact_bytes(content: bytes, *, filename: str, mime_type: str) -> ExtractedDocument:
    lower_name = filename.lower()
    if lower_name.endswith((".srt", ".vtt")) or mime_type in {"text/vtt", "application/x-subrip"}:
        text = content.decode("utf-8", errors="replace")
        if len(text) > _MAX_EXTRACTED_CHARACTERS:
            raise KnowledgebaseExtractionLimitExceeded("extracted text exceeds character limit")
        if text.count("-->") > _MAX_SOURCE_SPANS:
            raise KnowledgebaseExtractionLimitExceeded("document exceeds source span limit")
        document = _extract_transcript(text)
        if len(document.spans) > _MAX_SOURCE_SPANS:
            raise KnowledgebaseExtractionLimitExceeded("document exceeds source span limit")
        return document

    if mime_type.startswith(_TEXT_MIME_PREFIXES) or mime_type in _TEXT_MIME_TYPES:
        text = content.decode("utf-8", errors="replace")
        if len(text) > _MAX_EXTRACTED_CHARACTERS:
            raise KnowledgebaseExtractionLimitExceeded("extracted text exceeds character limit")
        if text.count("\n") + 1 > _MAX_SOURCE_SPANS:
            raise KnowledgebaseExtractionLimitExceeded("document exceeds source span limit")
        document_metadata: dict[str, Any] = {}
        extraction_method = "text"
        if mime_type == "text/markdown" or lower_name.endswith((".md", ".markdown", ".mdown")):
            text, document_metadata = _extract_markdown_frontmatter(text)
            if document_metadata:
                extraction_method = "markdown_frontmatter_v1"
        lines = text.splitlines(keepends=True)
        if len(lines) > _MAX_SOURCE_SPANS:
            raise KnowledgebaseExtractionLimitExceeded("document exceeds source span limit")
        spans: list[SourceSpan] = []
        char_pos = 0
        for line_no, line in enumerate(lines, start=1):
            line_text = line.rstrip("\r\n")
            spans.append(
                SourceSpan(
                    text=line_text,
                    locator={
                        "line_start": line_no,
                        "line_end": line_no,
                        "char_start": char_pos,
                        "char_end": char_pos + len(line_text),
                    },
                )
            )
            char_pos += len(line)
        return ExtractedDocument(
            spans=spans,
            extraction_method=extraction_method,
            metadata=document_metadata,
        )

    if mime_type == "application/pdf" or lower_name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - dependency is project-default
            raise RuntimeError("pypdf is required for PDF knowledgebase indexing") from exc
        reader = PdfReader(io.BytesIO(content))
        if len(reader.pages) > _MAX_PDF_PAGES:
            raise KnowledgebaseExtractionLimitExceeded("PDF exceeds page limit")
        spans = []
        extracted_characters = 0
        for page_index, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            extracted_characters += len(page_text)
            if extracted_characters > _MAX_EXTRACTED_CHARACTERS:
                raise KnowledgebaseExtractionLimitExceeded("extracted text exceeds character limit")
            if page_text.strip():
                spans.append(
                    SourceSpan(
                        text=page_text,
                        locator={"page_start": page_index, "page_end": page_index},
                    )
                )
        return ExtractedDocument(spans=spans, extraction_method="pdf")

    if lower_name.endswith(".docx") or mime_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            entries = archive.infolist()
            if len(entries) > _MAX_DOCX_ARCHIVE_ENTRIES:
                raise KnowledgebaseExtractionLimitExceeded("DOCX archive exceeds entry limit")
            if sum(entry.file_size for entry in entries) > _MAX_DOCX_UNCOMPRESSED_BYTES:
                raise KnowledgebaseExtractionLimitExceeded(
                    "DOCX expanded content exceeds size limit"
                )
            for entry in entries:
                if entry.flag_bits & 0x1:
                    raise KnowledgebaseExtractionLimitExceeded(
                        "encrypted DOCX archives are not supported"
                    )
                if entry.file_size > _MAX_DOCX_ENTRY_BYTES:
                    raise KnowledgebaseExtractionLimitExceeded(
                        "DOCX archive entry exceeds size limit"
                    )
                if entry.file_size / max(entry.compress_size, 1) > _MAX_DOCX_COMPRESSION_RATIO:
                    raise KnowledgebaseExtractionLimitExceeded(
                        "DOCX archive compression ratio exceeds limit"
                    )
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required for DOCX knowledgebase indexing") from exc
        document = Document(io.BytesIO(content))
        if len(document.paragraphs) > _MAX_SOURCE_SPANS:
            raise KnowledgebaseExtractionLimitExceeded("document exceeds source span limit")
        extracted_characters = 0
        spans: list[SourceSpan] = []
        for index, paragraph in enumerate(document.paragraphs, start=1):
            if not paragraph.text.strip():
                continue
            extracted_characters += len(paragraph.text)
            if extracted_characters > _MAX_EXTRACTED_CHARACTERS:
                raise KnowledgebaseExtractionLimitExceeded("extracted text exceeds character limit")
            spans.append(
                SourceSpan(
                    text=paragraph.text,
                    locator={"paragraph_start": index, "paragraph_end": index},
                )
            )
        return ExtractedDocument(spans=spans, extraction_method="docx")

    raise RuntimeError(f"Unsupported artifact type for knowledgebase indexing: {mime_type}")
